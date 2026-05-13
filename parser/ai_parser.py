import os
import json
import logging
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()
api_key = os.getenv("OPENROUTER_API_KEY")

log = logging.getLogger("doctoexam")

# OpenRouter connection
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=api_key
)

# Your strict list of exact subject spellings
ALLOWED_SUBJECTS = [
    "English", "Nepali", "Science", "Mathematics", "Samajik", "Computer", 
    "Opt.Math", "Abacus", "Pathfinder English", "Symphony", "G.K", 
    "Drawing", "All in One", "Sero-Fero", "Grammar", "Health", "Ltc.English"
]

def extract_text_from_docx(file_path: str) -> str:
    """Extracts all text, including tables, from the Word document."""
    doc = Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
                    
    return "\n".join(full_text)

def parse_document_with_ai(filepath: str, defaults: dict) -> tuple[dict, list]:
    """Reads doc, sends Filename and Text to AI, and forces exact subject matching."""
    
    raw_text = extract_text_from_docx(filepath)
    filename = os.path.basename(filepath) # Gets just the name, e.g. "Class 8 Grammar.docx"
    
    subjects_str = ", ".join(ALLOWED_SUBJECTS)
    
    system_prompt = f"""
    You are an expert data extractor. Extract the exam header information and multiple-choice questions.
    The document may be in English or Nepali.
    
    RULES FOR EXAM INFO (CRITICAL):
    1. Extract the "class" and "subject" primarily from the FILENAME provided. If missing in filename, look at the first few lines of text.
    2. The "subject" MUST EXACTLY MATCH one of these options: {subjects_str}.
       (Example: if filename says 'maths', output 'Mathematics'. If it says 'Grade 8 Science', output 'Science').
    3. Make the "chapter" EXACTLY the same string as the "subject".
    4. Extract the date if present, otherwise output "".
    5. There must be exactly 4 options per question. If you cannot find 4 options, set "needs_review" to true for that question.if there are more or less option then generate a similar answer that is wrong and set "needs_review" to true.
    
    RULES FOR QUESTIONS:
    1. Identify each question, its 4 options, and the correct answer.
    2. Correct answers are usually marked with a tick (✓), asterisk (*), or bold text. Remove the tick/asterisk from the clean option text.
    3. Options might be numbered a/b/c/d, 1/2/3/4, or क/ख/ग/घ. Map them STRICTLY to keys: "i", "ii", "iii", "iv".
    4. If you cannot find a correct answer for a question, set "needs_review" to true, "confidence" to "NONE", and "correct_option" to null.
    5. Default points per question is {defaults.get('points', 5)}.

    OUTPUT STRICTLY AS JSON MATCHING THIS EXACT SCHEMA:
    {{
      "exam_info": {{
        "class": "8",
        "subject": "Mathematics",
        "date": "083-01-26",
        "chapter": "Mathematics"
      }},
      "questions": [
        {{
          "number": 1,
          "question_raw": "Original question text",
          "question_clean": "Cleaned question text",
          "options": {{
            "i": "Option 1 clean text",
            "ii": "Option 2 clean text",
            "iii": "Option 3 clean text",
            "iv": "Option 4 clean text"
          }},
          "correct_option": "ii",
          "correct_text": "Option 2 clean text",
          "detection_method": "ai_detected",
          "confidence": "HIGH",
          "needs_review": false,
          "points": {defaults.get('points', 5)}
        }}
      ]
    }}
    """

    try:
        response = client.chat.completions.create(
            model="openai/gpt-4o-mini",
            response_format={ "type": "json_object" },
            temperature=0.0, # Zero creativity, strict adherence to rules
            messages=[
                {"role": "system", "content": system_prompt},
                # We feed the AI both the filename AND the raw text
                {"role": "user", "content": f"FILENAME: {filename}\n\nDOCUMENT TEXT:\n{raw_text}"}
            ]
        )
        
        parsed_data = json.loads(response.choices[0].message.content)
        
        exam_info = parsed_data.get("exam_info", {})
        questions = parsed_data.get("questions", [])
        
        return exam_info, questions

    except Exception as e:
        log.error(f"AI Parsing failed: {e}")
        return {}, []